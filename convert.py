import docx
import csv

# Extracting the field names encased within "$$" from the template file
def get_fields(template_doc_path):
    doc = docx.Document(template_doc_path)
    fields = []
    for para in doc.paragraphs:
        if '$' in para.text:
            start_idx = para.text.find('$')
            end_idx = para.text.rfind('$')
            if start_idx != -1 and end_idx != -1 and start_idx != end_idx:
                fields = para.text[start_idx:end_idx + 1]
                fields.append(fields)
    return fields

# Extracting the field values from the Word document for the fields encased within "$$"
def get_values(doc_path, fields):
    doc = docx.Document(doc_path)
    values = {field: None for field in fields}
    for para in doc.paragraphs:
        for field in fields:
            if field in para.text:
                values[field] = para.text.replace(field, "").strip()
    return values

# Inputting local paths to the template and the document files in question
template_doc_path = '/Users/soumyarajchatterjee/MS-Word template file.docx'
doc_path = '/Users/soumyarajchatterjee/MS-Word file.docx'

# Extracting fields from the template file
fields = get_fields(template_doc_path)
print(f"Extracted fields: {fields}")

# Extracting field values from the document file
values = get_values(doc_path, fields)
print(f"Extracted values: {values}")

# Writing final results to an output CSV file
csv_file_path = 'output.csv'
with open(csv_file_path, mode='w', newline='', encoding='utf-8') as file:
    writer = csv.writer(file)
    writer.writerow(['Field', 'Value'])
    for field, value in values.items():
        writer.writerow([field, value])

print(f"CSV file created: {csv_file_path}")






